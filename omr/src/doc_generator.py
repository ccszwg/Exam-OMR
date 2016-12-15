import os

import comtypes.client
import docx
import qrcode
from PyPDF2 import PdfFileMerger, PdfFileReader
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# todo: TURN INTO CLASS


def scrub(table_name):
    try:
        return int(table_name)
    except ValueError:
        if isinstance(table_name, str):
            return ''.join(chr for chr in table_name if chr.isalnum())
        else:
            return table_name


def generate_questions(num_questions, num_options):

    alphabet = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    text = ""

    for i in range(1, num_questions+1):
        text += "Question %s: \n   " % (str(i))

        for j in alphabet[0:num_options]:
            text += "%s   ◯      " % j

        if i == 10:
            text += "\n\n\n"
        else:
            text += "\n\n"

    return text


def replace_string(doc, holder, newtext):

    for paragraph in doc.paragraphs:
        if holder in paragraph.text:
            paragraph.text = newtext

    return doc


def activate_replacement(template, name, ID, num_questions, num_options, filelocation, class_ID, test_ID):
    try:
        doc = docx.Document(template)

        replace_string(doc, "{{IDENTITY}}", name)
        replace_string(doc, "{{QUESTIONS}}", generate_questions(num_questions, num_options))

        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=3,
            border=0,
        )

        print((str(ID) + " ") * 3 + (str(test_ID) + " ") * 3 + (str(class_ID) + " ") * 3)
        qr.add_data((str(ID) + " ") * 3 + (str(test_ID) + " ") * 3 + (str(class_ID) + " ") * 3)
        qr.make(fit=True)
        im = qr.make_image()

        im.save(filelocation + "/QRCODE_" + name + ".png")

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        r = p.add_run()
        r.add_picture(filelocation + "/QRCODE_" + name + ".png")

        return doc

    except Exception as e:
        print(e)


def covx_to_pdf(file_location):
    try:
        word = comtypes.client.CreateObject('Word.Application')

        for filename in os.listdir(file_location):
            if filename.endswith(".docx"):
                outfile = file_location + "/" + os.path.splitext(filename)[0] + ".pdf"
                infile = file_location + "/" + filename
                doc = word.Documents.Open(infile)
                doc.SaveAs(outfile, FileFormat=17)
                doc.Close()

    except Exception as e:
        print(e)
    finally:
        word.Quit()


def merge_pdfs(file_location):
    merger = PdfFileMerger()

    for filename in os.listdir(file_location):
        if filename.endswith(".pdf") and filename != "exam papers.pdf":
            with open(file_location + "/" + filename, "rb") as f:
                merger.append(PdfFileReader(f))

    merger.write(file_location + "/exam papers.pdf")

    for filename in os.listdir(file_location):
        if not filename == "exam papers.pdf":
            os.remove(file_location + "/" + filename)


def generate(names, num_questions, num_options, filelocation, class_ID, test_ID):

    for i in names:
        doc = activate_replacement("resources/template.docx", scrub(i["Name"]), int(i["ID"]), num_questions,
                                   num_options, filelocation, class_ID, test_ID)

        doc.save(filelocation + "/" + scrub(i["Name"]) + ".docx")

    covx_to_pdf(filelocation)
    merge_pdfs(filelocation)

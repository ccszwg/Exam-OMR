import os

import comtypes.client
import docx
from PyPDF2 import PdfFileMerger, PdfFileReader


def scrub(table_name):
    try:
        return int(table_name)
    except ValueError:
        if isinstance(table_name, str):
            return ''.join(chr for chr in table_name if chr.isalnum())
        else:
            return table_name


def generate_questions(num_questions, num_options):

    alphabet = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    text = ""

    for i in range(1, num_questions+1):
        text += "Question %s: \n   " % (str(i))

        for j in alphabet[0:num_options]:
            text += "%s   ◯      " % j

        if i == 10:
            text += "\n\n\n"
        else:
            text += "\n\n"

    return text


def replace_string(doc, holder, newtext):

    for paragraph in doc.paragraphs:
        if holder in paragraph.text:
            paragraph.text = newtext

    return doc


def replace_barcode(doc, barcode):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "{{BARCODE}}" in paragraph.text:
                        paragraph.text = barcode


def generate_barcode(ID):
    binary_ID = bin(ID)[2:].zfill(16)

    barcode = ""

    for i in binary_ID:
        if i == "1":
            barcode += " █ "
        else:
            barcode += "   "

    return barcode


def activate_replacement(template, name, ID, num_questions, num_options):

    doc = docx.Document(template)

    replace_string(doc, "{{IDENTITY}}", name)
    replace_string(doc, "{{QUESTIONS}}", generate_questions(num_questions, num_options))
    replace_barcode(doc, generate_barcode(ID))

    return doc


def covx_to_pdf(file_location):
    try:
        word = comtypes.client.CreateObject('Word.Application')

        for filename in os.listdir(file_location):
            if filename.endswith(".docx"):
                outfile = file_location + "/" + os.path.splitext(filename)[0] + ".pdf"
                infile = file_location + "/" + filename
                doc = word.Documents.Open(infile)
                doc.SaveAs(outfile, FileFormat=17)
                doc.Close()

    except Exception as e:
        print(e)
    finally:
        word.Quit()


def merge_pdfs(file_location):
    merger = PdfFileMerger()

    for filename in os.listdir(file_location):
        if filename.endswith(".pdf") and filename != "exam papers.pdf":
            with open(file_location + "/" + filename, "rb") as f:
                merger.append(PdfFileReader(f))

    merger.write(file_location + "/exam papers.pdf")

    for filename in os.listdir(file_location):
        if not filename == "exam papers.pdf":
            os.remove(file_location + "/" + filename)


def generate(names, num_questions, num_options, filelocation):

    for i in names:
        doc = activate_replacement("resources/template.docx", scrub(i["Name"]), int(i["ID"]), num_questions,
                                   num_options)

        doc.save(filelocation + "/" + scrub(i["Name"]) + ".docx")

    covx_to_pdf(filelocation)
    merge_pdfs(filelocation)
